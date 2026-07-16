from pathlib import Path
import docx
from docx.opc.exceptions import PackageNotFoundError
from ingestion.models import RawDocument
from ingestion.loaders.base import BaseLoader


class DOCXLoader(BaseLoader):
    """
    Loader for Microsoft Word (.docx) files.
    Extracts raw paragraph text (with style preservation) and table data, 
    while pulling native document metadata and statistics.
    """

    @property
    def name(self) -> str:
        return "docx"

    def load(self) -> RawDocument:
        path = Path(self.source)
        
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        if path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file, got '{path.suffix}'")

        full_text = []
        
        metadata = {
            "filename": path.name,
            "file_path": str(path),
            "file_size_bytes": path.stat().st_size,
            "loader": self.name,
            "paragraph_count": 0,
            "table_count": 0,
            "image_count": 0
        }

        try:
            doc = docx.Document(self.source)
        except PackageNotFoundError:
            raise ValueError(f"File is not a valid DOCX package or is corrupted: {path}")
        except Exception as e:
            raise RuntimeError(f"Fatal error loading DOCX file: {str(e)}")

        try:
            core_props = doc.core_properties
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.created:
                metadata["created_at"] = core_props.created.isoformat()
        except Exception as e:
            self.log_warning(f"Could not extract native core properties: {str(e)}")

        try:
            inline_shapes = doc.inline_shapes
            if len(inline_shapes) > 0:
                metadata["image_count"] = len(inline_shapes)
                self.log_warning(f"Document contains {len(inline_shapes)} images that were not processed.")
        except Exception as e:
            self.log_warning(f"Could not check for images: {str(e)}")

      
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                metadata["paragraph_count"] += 1
                style_name = para.style.name if para.style else "Normal"
                
               
                if style_name != "Normal":
                    full_text.append(f"[{style_name}] {text}")
                else:
                    full_text.append(text)

       
        for table in doc.tables:
            metadata["table_count"] += 1
            table_data = []
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_data.append(" | ".join(row_data))
            
            if table_data:
                full_text.append("\n".join(table_data))
               
                full_text.append("") 

       
        content = "\n\n".join(full_text).strip()

      
        if not content:
            self.log_warning("No textual content extracted from document.")

      
        metadata["character_count"] = len(content)
        metadata["word_count"] = len(content.split())

        return RawDocument(
            content=content,
            source=self.name,
            metadata=metadata,
            warnings=self.warnings
        )